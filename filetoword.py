# encoding: utf-8

"""

@version: 1.0

@author: HeKazhou

@file: docx_input

@time: 2020/8/10 10:54

"""

from docx import Document
from Abnormal import abnormal
import time
import pandas as pd
from docx.oxml.ns import  qn
from Abnormal.data import data
from Abnormal.abnormal import helper
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import shared
from pandas.api.types import is_string_dtype
from pandas.api.types import is_numeric_dtype

__author__ = 'HeKazhou'

class filetoword(object):




    def get_table(data,document,n):

        table=document.add_table(rows=len(data.index),cols=len(data.columns)+n, style='Table Grid')
        table.add_row()
        #添加表头
        for i in range(len(data.columns)):
            table.cell(0,i+n).text=data.columns[i]
        #添加表格
        for row in range(0,len(data.index)):
            for col in range(0,len(data.columns)):
                if n==0:
                    table.cell(row + 1,col).text=str(data.iloc[row,col])
                elif n==1:
                    table.cell(row + 1, 0).text = str(data.index[row])
                    table.cell(row + 1, col + 1).text = str(data.iloc[row, col])

        table.autofit = True
        return table

    def split_list(a_list):
        half = len(a_list)//2
        return a_list[:half], a_list[half:]

    def get_data_distribution(df,document,filename):
        data=[]
        str1=[]
        for c in range(len(df.columns)):
            if is_numeric_dtype(df.iloc[:, c]):
                data.append(df.columns[c])
            elif is_string_dtype(df.iloc[:, c]):
                str1.append(df.columns[c])
        for c in range(len(data)):
            if df[data[c]].count()!=0:
                r_zj = helper.get_count(df[data[c]],filename,data[c])
                r_zj = r_zj.reset_index()
                r_zj1 = r_zj.rename(columns={'index': data[c]})
                table = filetoword.get_table(r_zj1, document, 0)
                document.add_paragraph()
                document.add_picture('%s%s.png'%(filename,data[c]),width=shared.Inches(6.4),height=shared.Cm(8.6))
        for c in range(len(str1)):
            if df[str1[c]].count() != 0:
                s_zj = helper.get_count(df[str1[c]], filename,str1[c])
                s_zj = s_zj.reset_index()
                s_zj1 = s_zj.rename(columns={'index': str1[c]})
                table = filetoword.get_table(s_zj1, document, 0)
                document.add_paragraph()
                document.add_picture('%s%s.png'%(filename,str1[c]), width=shared.Inches(6.4), height=shared.Cm(8.6))
        return data,str1
    def filetoword(df,filename):
        data = []
        str1 = []
        list1=[]
        document = Document()
        style = document.styles['Normal']
        document.styles['Normal'].font.name = u'Times New Roman' #设置西文字体
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') #设置中文字体使用字体2->宋体

        run = document.add_heading('', 0).add_run(u"%s质量分析报告"%filename)
        run.font.name=u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.bold = True

        document.add_paragraph('创建时间：%s'%time.strftime('%Y-%m-%d'))

        run1 = document.add_heading('',1).add_run(u"数据概括")
        run1.font.name=u'宋体'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        document.add_paragraph('数据来源', style='List Bullet')

        first_time,last_time=helper.get_time(df,'ETL_DATE')
        # 首行缩进两个字符
        #paragraph_format = style.paragraph_format
        document.add_paragraph('%s -- %s入库的数据'%(first_time,last_time)).paragraph_format.left_indent = Cm(0.75)
        #document.add_paragraph('%s -- %s入库的数据'%(first_time,last_time))

        document.add_paragraph('总记录数（反映数据规模）', style='List Bullet')

        count_sum = helper.count_sum(df, 'ETL_DATE')

        table=filetoword.get_table(count_sum,document,0)

        document.add_paragraph()




        document.add_paragraph('分年份记录数（柱状图）', style='List Bullet')

        count_year=helper.count_plot_year(df,filename,'ETL_DATE')

        document.add_picture('%s时间分布图.png' % filename, width=shared.Inches(6.4), height=shared.Cm(8.6))

        table=filetoword.get_table(count_year,document,1)

        run2 = document.add_heading('',1).add_run(u"值分析")

        document.add_paragraph('空值占比（查询每个字段的缺失百分比）', style='List Bullet')

        merge_count=helper.merge_save_count(df)

        table=filetoword.get_table(merge_count,document,1)


        document.add_paragraph('异常值占比（查询每个字段的异常值百分比）', style='List Bullet')


        document.add_paragraph('变量本身的分布', style='List Bullet')

        data,str1=filetoword.get_data_distribution(df,document,filename)

        # for c in range(len(df.columns)):
        #     if is_numeric_dtype(df.iloc[:, c]):
        #         data.append(df.columns[c])
        #     elif is_string_dtype(df.iloc[:, c]):
        #         str1.append(df.columns[c])
        # for c in range(len(data)):
        #     if df[data[c]].count()!=0:
        #         r_zj = helper.get_count(df[data[c]],data[c])
        #         r_zj = r_zj.reset_index()
        #         r_zj1 = r_zj.rename(columns={'index': data[c]})
        #         table = filetoword.get_table(r_zj1, document, 0)
        #         document.add_paragraph()
        #         document.add_picture('%s%s.png'%(name,data[c]),width=shared.Inches(6.4),height=shared.Cm(8.6))
        # for c in range(len(str1)):
        #     if df[str1[c]].count() != 0:
        #         s_zj = helper.get_count(df[str1[c]], str1[c])
        #         s_zj = s_zj.reset_index()
        #         s_zj1 = s_zj.rename(columns={'index': str1[c]})
        #         table = filetoword.get_table(s_zj1, document, 0)
        #         document.add_paragraph()
        #         document.add_picture('%s%s.png'%(name,str1[c]), width=shared.Inches(6.4), height=shared.Cm(8.6))


        # for i in range(len(data_distribution)):
        #     if len(data_distribution[i]) == 0:
        #         continue
        #     A,B=split_list(data_distribution[i])
        #     A=pd.DataFrame(A,columns=Name[i])
        #     B=pd.DataFrame(B,columns='counts')
        #     df = pd.merge(A, B, how='inner', left_index=True, right_index=True)
        #     table = get_table(df, document, 0)



        #document.add_picture()
        #document.add_paragraph('变量间关系分布', style='List Bullet')

        run3 = document.add_heading('',1).add_run(u"统计分析")
        document.add_paragraph('众数（发生频率最高的值，当异常值出现频率最高，则需要考虑数据可靠性）', style='List Bullet')
        document.add_paragraph('五数概括法（minimum，25% percentile，median，75% percentile，maximum）及均值（数据平均状况）', style='List Bullet')
        document.add_paragraph('偏度（检验分布正态性，对于符合正态分布的字段可采取3σ原则判断异常值）', style='List Bullet')
        document.add_paragraph('标准差（数据分散程度）', style='List Bullet')

        for c in data:
            if df[c].count() != 0:
                list1.append(c)
        df1=df.loc[:,list1]
        #statistic_analysis=helper.statistic_analysis(df)
        #table=get_table(statistic_analysis,document,1)
        # cond = df[(df.loc[:,data]).count() !=0]
        # print(cond)

        statistic_analysis = helper.statistic_analysis(df1)
        table = filetoword.get_table(statistic_analysis, document, 1)
        document.add_paragraph()

        document.save('demo.docx')


if __name__ == '__main__':
    data=helper.read_data(r'./op.xlsx')
    filetoword.filetoword(data,'门诊表')